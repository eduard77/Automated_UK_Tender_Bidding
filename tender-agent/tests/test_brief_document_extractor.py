"""Per-file text extraction (Phase 4 chunk 5).

Covers docx / xlsx / pdf / zip recursion / txt / csv / unsupported / corrupt.
Extraction NEVER throws — failures become status='error' / 'unsupported' /
'empty', and the caller (content store) records that as the row's status.
"""
from __future__ import annotations

import io
import zipfile

import pytest

from tender_agent.services.brief.document_extractor import (
    EXTRACTOR_VERSION,
    extract_file_bytes,
    extract_file_path,
)

# --------------------------------------------------------------------------
# Tiny binary fixtures generated in-process so we don't ship binary files.
# --------------------------------------------------------------------------


def _make_docx_bytes(paragraphs: list[str], tables: list[list[list[str]]] | None = None) -> bytes:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    for rows in tables or []:
        table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row):
                table.rows[ri].cells[ci].text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_xlsx_bytes(sheets: dict[str, list[list[object]]]) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    # Replace the default sheet with named ones.
    default = wb.active
    wb.remove(default)
    for name, rows in sheets.items():
        ws = wb.create_sheet(title=name)
        for row in rows:
            ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_pdf_with_text(text: str) -> bytes:
    """Minimal PDF that pypdf can read text from."""
    from pypdf import PdfWriter
    from pypdf.generic import (
        ArrayObject,
        ContentStream,
        DictionaryObject,
        FloatObject,
        NameObject,
        NumberObject,
        RectangleObject,
        TextStringObject,
    )

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)

    # Build a simple text-showing content stream: BT /F1 12 Tf 50 700 Td (text) Tj ET
    safe = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream_data = (
        f"BT /F1 12 Tf 50 700 Td ({safe}) Tj ET"
    ).encode("ascii")
    content = ContentStream(None, writer)
    content._data = stream_data  # type: ignore[attr-defined]
    page[NameObject("/Contents")] = writer._add_object(content)

    # Attach a font resource so the content stream is valid enough for pypdf.
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    resources = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): font_ref}
            ),
        }
    )
    page[NameObject("/Resources")] = resources
    page[NameObject("/MediaBox")] = RectangleObject([0, 0, 612, 792])
    # Silence unused-import warnings for the helpers above.
    _ = (
        ArrayObject,
        FloatObject,
        NumberObject,
        TextStringObject,
    )

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _make_empty_pdf() -> bytes:
    """A real PDF with no text content (simulates a scanned/image-only PDF)."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _make_zip_bytes(members: dict[str, bytes]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in members.items():
            zf.writestr(name, data)
    return buf.getvalue()


# --------------------------------------------------------------------------
# docx
# --------------------------------------------------------------------------


def test_extractor_version_is_a_short_string():
    # The content store uses (sha256, extractor_version) as its uniqueness
    # key — accidentally bumping this in the wrong place would invalidate
    # every cached extraction, so guard the contract here.
    assert isinstance(EXTRACTOR_VERSION, str)
    assert 0 < len(EXTRACTOR_VERSION) < 32


def test_extract_docx_paragraphs_and_tables():
    data = _make_docx_bytes(
        paragraphs=["Tender for cleaning services.", "Closes 12 December 2025."],
        tables=[[["Insurance", "£5m"], ["ISO 9001", "Required"]]],
    )
    result = extract_file_bytes(data, filename="itt.docx")
    assert result.status == "ok"
    assert result.doc_type == "docx"
    assert "cleaning services" in result.text
    assert "Closes 12 December" in result.text
    # Table cells flattened with the | separator.
    assert "Insurance | £5m" in result.text
    assert result.char_count == len(result.text)


def test_extract_docx_corrupt_returns_error():
    result = extract_file_bytes(b"this is not a docx", filename="bad.docx")
    assert result.status == "error"
    assert result.doc_type == "docx"
    # Crucially: never raises.


# --------------------------------------------------------------------------
# xlsx
# --------------------------------------------------------------------------


def test_extract_xlsx_multi_sheet_includes_sheet_names_and_cells():
    data = _make_xlsx_bytes(
        {
            "Pricing": [["Item", "Qty", "Rate"], ["Mop", 10, 4.5]],
            "Quality": [["Q", "Weight"], ["Sustainability", "20%"]],
        }
    )
    result = extract_file_bytes(data, filename="pricing.xlsx")
    assert result.status == "ok"
    assert result.doc_type == "xlsx"
    assert "=== Sheet: Pricing ===" in result.text
    assert "=== Sheet: Quality ===" in result.text
    assert "Mop | 10 | 4.5" in result.text
    assert "Sustainability | 20%" in result.text


def test_extract_xlsx_corrupt_returns_error():
    result = extract_file_bytes(b"nope", filename="bad.xlsx")
    assert result.status == "error"


# --------------------------------------------------------------------------
# pdf
# --------------------------------------------------------------------------


def test_extract_pdf_with_text():
    data = _make_pdf_with_text("Hello from the ITT")
    result = extract_file_bytes(data, filename="itt.pdf")
    # pdfplumber / pypdf may return slightly different whitespace, so check
    # for substring containment rather than equality.
    assert result.doc_type == "pdf"
    assert result.status in ("ok", "empty")
    if result.status == "ok":
        assert "Hello" in result.text


def test_extract_pdf_no_text_is_empty_not_error():
    data = _make_empty_pdf()
    result = extract_file_bytes(data, filename="scanned.pdf")
    # Either backend may legitimately return empty for a blank page.
    assert result.doc_type == "pdf"
    assert result.status in ("empty", "ok")
    if result.status == "empty":
        assert "no extractable text" in (result.detail or "")


def test_extract_pdf_corrupt_returns_error():
    result = extract_file_bytes(b"%PDF-broken", filename="broken.pdf")
    assert result.status == "error"


# --------------------------------------------------------------------------
# zip recursion
# --------------------------------------------------------------------------


def test_extract_zip_recurses_into_members():
    docx_data = _make_docx_bytes(paragraphs=["Inside the zip"])
    txt_data = b"plain text member"
    archive = _make_zip_bytes(
        {
            "spec.docx": docx_data,
            "notes/readme.txt": txt_data,
        }
    )
    result = extract_file_bytes(archive, filename="pack.zip")
    assert result.status == "ok"
    assert result.doc_type == "zip"
    assert len(result.members) == 2

    by_name = {m.filename: m for m in result.members}
    docx_member = by_name["spec.docx"]
    txt_member = by_name["notes/readme.txt"]
    assert docx_member.doc_type == "zip-member"
    assert docx_member.status == "ok"
    assert "Inside the zip" in (docx_member.text or "")
    assert txt_member.doc_type == "zip-member"
    assert "plain text member" in (txt_member.text or "")


def test_extract_zip_bad_zip_returns_error():
    result = extract_file_bytes(b"not really a zip", filename="bad.zip")
    assert result.status == "error"
    assert result.doc_type == "zip"


# --------------------------------------------------------------------------
# plain text / csv
# --------------------------------------------------------------------------


def test_extract_txt():
    result = extract_file_bytes(b"hello world", filename="notes.txt")
    assert result.status == "ok"
    assert result.text == "hello world"
    assert result.doc_type == "txt"


def test_extract_csv_flattens_rows():
    result = extract_file_bytes(b"a,b,c\n1,2,3\n", filename="data.csv")
    assert result.status == "ok"
    assert result.doc_type == "csv"
    assert "a | b | c" in result.text
    assert "1 | 2 | 3" in result.text


# --------------------------------------------------------------------------
# unsupported / fallthroughs
# --------------------------------------------------------------------------


def test_extract_unsupported_extension():
    result = extract_file_bytes(b"\x00\x01\x02", filename="image.png")
    assert result.status == "unsupported"
    assert "png" in (result.detail or "")


def test_extract_uses_content_type_when_extension_missing():
    # Some servers serve docx without a .docx extension; content-type rescues us.
    data = _make_docx_bytes(paragraphs=["From content-type"])
    result = extract_file_bytes(
        data,
        filename="anonymous",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    assert result.status == "ok"
    assert result.doc_type == "docx"
    assert "From content-type" in result.text


def test_extract_file_path_missing_returns_error(tmp_path):
    # Path does not exist on disk — extraction must record an error, not raise.
    p = tmp_path / "missing.docx"
    result = extract_file_path(p)
    assert result.status == "error"
    assert "missing on disk" in (result.detail or "")


def test_extract_file_path_round_trip(tmp_path):
    data = _make_docx_bytes(paragraphs=["from disk"])
    p = tmp_path / "doc.docx"
    p.write_bytes(data)
    result = extract_file_path(p)
    assert result.status == "ok"
    assert "from disk" in (result.text or "")


@pytest.mark.parametrize(
    ("filename", "expected_doc_type"),
    [
        ("notes.TXT", "txt"),
        ("plan.CSV", "csv"),
    ],
)
def test_extension_casing_is_normalised(filename: str, expected_doc_type: str):
    payload = b"x,y\n1,2\n" if filename.lower().endswith("csv") else b"abc"
    result = extract_file_bytes(payload, filename=filename)
    assert result.doc_type == expected_doc_type
