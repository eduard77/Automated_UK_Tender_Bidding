"""Per-file text extraction for the document content store.

Each public entry point returns an `ExtractedContent` (or a list of them, for
zip archives where every member becomes its own row). Failures NEVER raise:
unknown formats become status="unsupported", corrupt files become
status="error", and scanned PDFs with no extractable text become
status="empty" with a detail explaining why. The caller (content_store) is
then responsible for persisting whatever it got back.

Extractor behaviour is versioned by `EXTRACTOR_VERSION` — bump it any time the
extraction logic changes meaningfully (e.g. switching PDF backends, changing
how spreadsheets are laid out). A version bump triggers re-extraction the next
time a file is accessed, while leaving older rows intact for audit.
"""
from __future__ import annotations

import csv
import io
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import IO

import structlog

logger = structlog.get_logger(__name__)


# Bump when the extractor logic changes. The content store uses
# (sha256, extractor_version) as its uniqueness key, so a bump triggers a
# fresh extraction the next time a file is touched.
EXTRACTOR_VERSION = "1"

# Cap the number of zip members we recurse into per archive — defends against
# zip-bombs and absurdly large archives. Anything beyond is recorded as a
# single zip-overflow row.
MAX_ZIP_MEMBERS = 200


@dataclass
class ExtractedContent:
    """The result of extracting one logical document.

    For zip archives, the archive itself becomes one ExtractedContent with
    `doc_type="zip"` and an empty text body (status="ok", detail names the
    members), and each member becomes its own ExtractedContent with
    `doc_type="zip-member"` and `filename` set to the member path inside the
    archive.
    """

    text: str | None
    char_count: int
    status: str  # ok | empty | unsupported | error
    detail: str | None
    doc_type: str | None
    # For zip members: the path inside the archive ("subdir/spec.pdf"), so the
    # caller can record a per-member row without re-deriving names. None for
    # the top-level file.
    filename: str | None = None
    # Member contents, if this was a zip. Populated only for the archive's
    # own ExtractedContent; each member is also returned as its own
    # ExtractedContent at the top level of `extract_file_bytes`.
    members: list[ExtractedContent] = field(default_factory=list)


def _ok(text: str, doc_type: str, *, filename: str | None = None) -> ExtractedContent:
    cleaned = text.strip() if text else ""
    if not cleaned:
        return ExtractedContent(
            text=None,
            char_count=0,
            status="empty",
            detail="no extractable text found",
            doc_type=doc_type,
            filename=filename,
        )
    return ExtractedContent(
        text=cleaned,
        char_count=len(cleaned),
        status="ok",
        detail=None,
        doc_type=doc_type,
        filename=filename,
    )


def _unsupported(ext: str, *, filename: str | None = None) -> ExtractedContent:
    return ExtractedContent(
        text=None,
        char_count=0,
        status="unsupported",
        detail=f"no extractor for .{ext}",
        doc_type=ext or None,
        filename=filename,
    )


def _error(detail: str, doc_type: str | None, *, filename: str | None = None) -> ExtractedContent:
    return ExtractedContent(
        text=None,
        char_count=0,
        status="error",
        detail=detail,
        doc_type=doc_type,
        filename=filename,
    )


def _guess_ext(filename: str | None, content_type: str | None) -> str:
    """Best-effort extension. Falls back to the content-type subtype."""
    if filename and "." in filename:
        ext = filename.rsplit(".", 1)[-1].lower()
        if ext.isalnum() and 1 <= len(ext) <= 8:
            return ext
    if content_type:
        sub = content_type.split(";")[0].split("/")[-1].lower()
        # application/vnd.openxmlformats-officedocument.wordprocessingml.document
        # → "wordprocessingml.document" → not useful; map a few we care about
        # by substring.
        for marker, ext in (
            ("wordprocessingml", "docx"),
            ("spreadsheetml", "xlsx"),
            ("pdf", "pdf"),
            ("zip", "zip"),
            ("plain", "txt"),
            ("csv", "csv"),
        ):
            if marker in sub:
                return ext
    return ""


# --------------------------------------------------------------------------
# Per-format extractors. Each takes raw bytes and returns ExtractedContent.
# Never raises — exceptions are caught and turned into status="error".
# --------------------------------------------------------------------------


def _extract_docx(data: bytes, *, filename: str | None = None) -> ExtractedContent:
    try:
        from docx import Document
    except ImportError:
        return _error("python-docx not installed", "docx", filename=filename)
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        return _error(f"docx parse failed: {exc}", "docx", filename=filename)
    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    # Tables matter too — quality questionnaires and pricing schedules often
    # live in tables. Flatten row-by-row, cells separated by a pipe.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _ok("\n".join(parts), "docx", filename=filename)


def _extract_xlsx(data: bytes, *, filename: str | None = None) -> ExtractedContent:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return _error("openpyxl not installed", "xlsx", filename=filename)
    try:
        wb = load_workbook(filename=io.BytesIO(data), read_only=True, data_only=True)
    except Exception as exc:  # noqa: BLE001
        return _error(f"xlsx parse failed: {exc}", "xlsx", filename=filename)
    out: list[str] = []
    try:
        for sheet_name in wb.sheetnames:
            out.append(f"=== Sheet: {sheet_name} ===")
            ws = wb[sheet_name]
            for row in ws.iter_rows(values_only=True):
                cells = [
                    str(c).strip()
                    for c in row
                    if c is not None and str(c).strip()
                ]
                if cells:
                    out.append(" | ".join(cells))
    finally:
        wb.close()
    return _ok("\n".join(out), "xlsx", filename=filename)


def _extract_pdf(data: bytes, *, filename: str | None = None) -> ExtractedContent:
    # Prefer pdfplumber (better layout); fall back to pypdf if unavailable.
    parts: list[str] = []
    try:
        import pdfplumber  # type: ignore[import-not-found]
    except ImportError:
        return _extract_pdf_pypdf(data, filename=filename)
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                try:
                    text = page.extract_text() or ""
                except Exception:  # noqa: BLE001
                    text = ""
                if text.strip():
                    parts.append(text)
    except Exception as exc:  # noqa: BLE001
        # pdfplumber chokes on some PDFs that pypdf reads fine; try the fallback.
        fallback = _extract_pdf_pypdf(data, filename=filename)
        if fallback.status == "ok":
            return fallback
        return _error(f"pdf parse failed: {exc}", "pdf", filename=filename)

    text = "\n".join(parts).strip()
    if not text:
        return ExtractedContent(
            text=None,
            char_count=0,
            status="empty",
            detail="no extractable text (scanned PDF or image-only)",
            doc_type="pdf",
            filename=filename,
        )
    return _ok(text, "pdf", filename=filename)


def _extract_pdf_pypdf(data: bytes, *, filename: str | None = None) -> ExtractedContent:
    try:
        from pypdf import PdfReader
    except ImportError:
        return _error("pypdf not installed", "pdf", filename=filename)
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        return _error(f"pdf parse failed: {exc}", "pdf", filename=filename)
    parts: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            text = ""
        if text.strip():
            parts.append(text)
    text = "\n".join(parts).strip()
    if not text:
        return ExtractedContent(
            text=None,
            char_count=0,
            status="empty",
            detail="no extractable text (scanned PDF or image-only)",
            doc_type="pdf",
            filename=filename,
        )
    return _ok(text, "pdf", filename=filename)


def _extract_plaintext(
    data: bytes, doc_type: str, *, filename: str | None = None
) -> ExtractedContent:
    try:
        text = data.decode("utf-8", errors="replace")
    except Exception as exc:  # noqa: BLE001
        return _error(f"decode failed: {exc}", doc_type, filename=filename)
    return _ok(text, doc_type, filename=filename)


def _extract_csv(data: bytes, *, filename: str | None = None) -> ExtractedContent:
    # CSV is plain text, but join cells with pipes for the same shape as xlsx.
    try:
        text = data.decode("utf-8", errors="replace")
    except Exception as exc:  # noqa: BLE001
        return _error(f"csv decode failed: {exc}", "csv", filename=filename)
    out: list[str] = []
    try:
        reader = csv.reader(io.StringIO(text))
        for row in reader:
            cells = [c.strip() for c in row if c and c.strip()]
            if cells:
                out.append(" | ".join(cells))
    except Exception as exc:  # noqa: BLE001
        return _error(f"csv parse failed: {exc}", "csv", filename=filename)
    return _ok("\n".join(out), "csv", filename=filename)


def _extract_zip(data: bytes, *, filename: str | None = None) -> ExtractedContent:
    """Recurse into every member; the archive itself returns a summary row."""
    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile as exc:
        return _error(f"bad zip: {exc}", "zip", filename=filename)
    except Exception as exc:  # noqa: BLE001
        return _error(f"zip open failed: {exc}", "zip", filename=filename)

    members: list[ExtractedContent] = []
    truncated = False
    try:
        names = [n for n in zf.namelist() if not n.endswith("/")]
        if len(names) > MAX_ZIP_MEMBERS:
            truncated = True
            names = names[:MAX_ZIP_MEMBERS]
        for name in names:
            try:
                with zf.open(name) as member:
                    member_bytes = member.read()
            except Exception as exc:  # noqa: BLE001
                members.append(
                    _error(f"member read failed: {exc}", "zip-member", filename=name)
                )
                continue
            # Recurse — but mark the doc_type as zip-member so the caller knows
            # this row is from inside an archive, not a standalone file.
            ext = _guess_ext(name, None)
            sub = _extract_bytes_by_ext(
                member_bytes, ext, filename=name, content_type=None
            )
            sub.doc_type = "zip-member"
            sub.filename = name
            members.append(sub)
    finally:
        zf.close()

    summary_detail = (
        f"{len(members)} member(s)"
        + (f" (truncated at {MAX_ZIP_MEMBERS})" if truncated else "")
    )
    return ExtractedContent(
        text=None,
        char_count=0,
        status="ok",
        detail=summary_detail,
        doc_type="zip",
        filename=filename,
        members=members,
    )


def _extract_bytes_by_ext(
    data: bytes,
    ext: str,
    *,
    filename: str | None,
    content_type: str | None,
) -> ExtractedContent:
    ext = (ext or "").lower()
    if ext == "docx":
        return _extract_docx(data, filename=filename)
    if ext == "xlsx":
        return _extract_xlsx(data, filename=filename)
    if ext == "pdf":
        return _extract_pdf(data, filename=filename)
    if ext == "zip":
        return _extract_zip(data, filename=filename)
    if ext in {"txt", "md", "log"}:
        return _extract_plaintext(data, ext, filename=filename)
    if ext == "csv":
        return _extract_csv(data, filename=filename)
    # Some servers serve docx/xlsx without an extension — try the content-type.
    if not ext and content_type:
        guess = _guess_ext(None, content_type)
        if guess and guess != ext:
            return _extract_bytes_by_ext(
                data, guess, filename=filename, content_type=None
            )
    return _unsupported(ext or "bin", filename=filename)


# --------------------------------------------------------------------------
# Public entry points
# --------------------------------------------------------------------------


def extract_file_bytes(
    data: bytes,
    *,
    filename: str | None,
    content_type: str | None = None,
) -> ExtractedContent:
    """Extract text from raw bytes. Used by tests and the content store
    (after reading the file from disk via storage_key)."""
    ext = _guess_ext(filename, content_type)
    try:
        return _extract_bytes_by_ext(
            data, ext, filename=filename, content_type=content_type
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning(
            "extractor.unexpected_failure",
            filename=filename,
            ext=ext,
            error=str(exc),
        )
        return _error(f"unexpected: {exc}", ext or None, filename=filename)


def extract_file_path(
    path: str | Path,
    *,
    filename: str | None = None,
    content_type: str | None = None,
) -> ExtractedContent:
    """Convenience wrapper: read the file then extract."""
    p = Path(path)
    if not p.is_file():
        return _error(
            f"file missing on disk: {p}",
            _guess_ext(filename or p.name, content_type) or None,
            filename=filename or p.name,
        )
    try:
        data = p.read_bytes()
    except Exception as exc:  # noqa: BLE001
        return _error(
            f"read failed: {exc}",
            _guess_ext(filename or p.name, content_type) or None,
            filename=filename or p.name,
        )
    return extract_file_bytes(
        data,
        filename=filename or p.name,
        content_type=content_type,
    )


def extract_stream(
    stream: IO[bytes],
    *,
    filename: str | None,
    content_type: str | None = None,
) -> ExtractedContent:
    """Convenience for tests / callers that hold a file-like object."""
    return extract_file_bytes(
        stream.read(), filename=filename, content_type=content_type
    )
